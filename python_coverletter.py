from docx import Document
from fpdf import FPDF
from datetime import date
import os
from pywebio.input import textarea, input, input_group
from pywebio.output import put_text, put_buttons

def get_current_date():
    today = date.today()
    return today.strftime("%B %d, %Y")

def generate_cover_letter(job_description, skill_set, position, company, user_name, address, city_state_zip, email, phone):
    # Create a new Document
    doc = Document()

    # Add user's name and position
    doc.add_paragraph(user_name).bold = True
    doc.add_paragraph(address).bold = True
    doc.add_paragraph(city_state_zip).bold = True
    doc.add_paragraph(email).bold = True
    doc.add_paragraph(phone).bold = True
    doc.add_paragraph(f'Date: {get_current_date()}').italic = True

    # Add recipient's details
    doc.add_paragraph(f'\n\nDear Hiring Manager,\n\n')

    # Add main content
    p = doc.add_paragraph()
    p.add_run(f'I am writing to express my strong interest in the position of {position} at {company}. ').bold = True
    p.add_run(f'With a background across the mentioned abilities that a candidate should have {job_description} and a skill set that includes {", ".join(skill_set)}, '
              f'I am self-assured in my capacity to contribute effectively to your team and meet the goals. I am very enthusiastic and excited to start my career at {company}\n')

    # Add closing statement
    doc.add_paragraph('Thank you for considering my application. I am looking forward to the opportunity to '
                      'discuss how my skills and experiences align with the requirements of the position.')

    # Add closing credits
    doc.add_paragraph(f'\nYours sincerely, \n\n{user_name}.')

    # Save the document
    doc.save('cover_letter.docx')

def convert_to_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)

    # Read the content from the docx file
    doc = Document('cover_letter.docx')
    for para in doc.paragraphs:
        pdf.multi_cell(0, 10, txt=para.text.encode('latin-1', 'replace').decode('latin-1'))

    # Get the path to the desktop
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    pdf_path = os.path.join(desktop_path, 'cover_letter.pdf')

    # Save the pdf to the desktop
    pdf.output(pdf_path)

def cover_letter_app():
    
    user_name = input("Enter your name:", placeholder="Your Name")
    skill_set = textarea("Enter your skill set:", rows=5, placeholder="Your Skill Set")
    address = textarea("Enter your address:", rows=5, placeholder="Your Address")
    city_state_zip = textarea("Enter you city, state and zip:", rows=5, placeholder="Your City State and Zip")
    email = input("Enter your email address:", placeholder="Your Email")
    phone = input("Enter your phone number:", placeholder="Your Phone Number")


    position = input("Enter offered position:", placeholder="Position")
    company = input("Enter company name:", placeholder="Company Name")
    job_description = textarea("Enter job description:", rows=5, placeholder="Job Description")
    
  

    generate_cover_letter(job_description, skill_set, position, company, user_name , address, city_state_zip, email, phone)
    convert_to_pdf()

    # Read the content from the docx file
    doc = Document('cover_letter.docx')
    cover_letter_content = '\n'.join([para.text for para in doc.paragraphs])
    
    # Display the cover letter content
    put_text(cover_letter_content)

if __name__ == '__main__':
    cover_letter_app()